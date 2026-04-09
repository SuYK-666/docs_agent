from __future__ import annotations

import argparse
import json
import logging
import re
from datetime import date, datetime
from html import escape as html_escape
from pathlib import Path
from typing import Any, Iterable, Mapping

from jinja2 import Environment, FileSystemLoader, select_autoescape
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from config.logger_setup import get_log_session, get_logger, log_step, setup_logger, to_relative_path


CURRENT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = CURRENT_DIR.parent
DEFAULT_TEMPLATE_DIR = CURRENT_DIR / "templates"
DEFAULT_CACHE_DIR = PROJECT_ROOT / "data_workspace" / "processed_cache"
DEFAULT_REPORT_DIR = PROJECT_ROOT / "data_workspace" / "final_reports" / "reports"
LOGGER = logging.getLogger("docs_agent.report_renderer")

TITLE_TEXT = "公文速阅报告报送"
TITLE_FONT = "方正小标宋简体"
SUBTITLE_FONT = "楷体_GB2312"
HEADING2_FONT = "黑体"
HEADING3_FONT = "楷体_GB2312"
BODY_FONT = "仿宋_GB2312"
LATIN_FONT = "Times New Roman"
COLOR_BLACK = RGBColor(0, 0, 0)
COLOR_RED = RGBColor(185, 28, 28)
SIZE_PT_LEVEL1 = 22
SIZE_PT_LEVEL2 = 16
SIZE_PT_LEVEL3 = 16
LINE_SPACING_PT = 28
BODY_FIRST_LINE_INDENT_PT = 32

PROVIDER_DISPLAY_NAME = {
	"deepseek": "DeepSeek",
	"tongyi": "阿里通义系列",
	"wenxin": "百度文心系列",
	"gaoding": "稿定设计",
	"modelwhale": "和鲸ModelWhale",
	"jimeng": "即梦",
	"doubao": "豆包AI",
	"spark": "科大讯飞星火",
	"kimi": "Kimi",
	"hunyuan": "腾讯混元系列",
	"zhipu": "智谱AI",
}
SUPPORTED_REPORT_FORMATS = {"md", "html", "docx"}


def _safe_text(value: Any, fallback: str = "未提及") -> str:
	text = str(value).strip() if value is not None else ""
	return text or fallback


def _summary_lines(summary: str) -> list[str]:
	parts = [line.strip() for line in str(summary).splitlines() if line.strip()]
	return parts or ["未提及"]


def _safe_filename(name: str) -> str:
	cleaned = re.sub(r"[\\/:*?\"<>|]+", "_", name).strip()
	return cleaned or "report"


def _build_word_report_filename(now: datetime | None = None) -> str:
	timestamp = (now or datetime.now()).strftime("%Y年%m月%d日%H时%M分%S秒")
	return f"公文速阅报告_{timestamp}"


def _build_word_timestamp_line(now: datetime | None = None) -> str:
	timestamp = (now or datetime.now()).strftime("%Y年%m月%d日%H时%M分%S秒")
	return f"（{timestamp}整理）"


def _chinese_index(index: int) -> str:
	digits = {
		0: "零",
		1: "一",
		2: "二",
		3: "三",
		4: "四",
		5: "五",
		6: "六",
		7: "七",
		8: "八",
		9: "九",
		10: "十",
	}
	if index <= 10:
		return digits.get(index, str(index))
	if index < 20:
		return f"十{digits.get(index - 10, '')}"
	if index < 100:
		tens = index // 10
		ones = index % 10
		tens_text = f"{digits.get(tens, str(tens))}十"
		return tens_text if ones == 0 else f"{tens_text}{digits.get(ones, str(ones))}"
	return str(index)


def _resolve_model_display(data: Mapping[str, Any]) -> str:
	provider_raw = _safe_text(
		data.get("llm_provider"),
		fallback=_safe_text(data.get("pipeline_meta", {}).get("llm_provider"), fallback="deepseek") if isinstance(data.get("pipeline_meta"), Mapping) else "deepseek",
	)
	provider_key = provider_raw.lower()
	if provider_key in PROVIDER_DISPLAY_NAME:
		return PROVIDER_DISPLAY_NAME[provider_key]
	if provider_raw:
		return provider_raw
	return "DeepSeek"


def _set_run_font(run: Any, east_asia_font: str, size_pt: int, bold: bool = False) -> None:
	run.font.bold = bold
	run.font.size = Pt(size_pt)
	run.font.color.rgb = COLOR_BLACK
	run.font.name = LATIN_FONT

	rpr = run._element.get_or_add_rPr()
	rfonts = rpr.rFonts
	if rfonts is None:
		rfonts = OxmlElement("w:rFonts")
		rpr.insert(0, rfonts)
	rfonts.set(qn("w:ascii"), LATIN_FONT)
	rfonts.set(qn("w:hAnsi"), LATIN_FONT)
	rfonts.set(qn("w:eastAsia"), east_asia_font)


def _set_style_font(style: Any, east_asia_font: str, size_pt: int) -> None:
	style.font.bold = False
	style.font.size = Pt(size_pt)
	style.font.color.rgb = COLOR_BLACK
	style.font.name = LATIN_FONT

	rpr = style.element.get_or_add_rPr()
	rfonts = rpr.rFonts
	if rfonts is None:
		rfonts = OxmlElement("w:rFonts")
		rpr.insert(0, rfonts)
	rfonts.set(qn("w:ascii"), LATIN_FONT)
	rfonts.set(qn("w:hAnsi"), LATIN_FONT)
	rfonts.set(qn("w:eastAsia"), east_asia_font)


def _get_or_create_paragraph_style(doc: Document, east_asia_font: str, size_pt: int) -> Any:
	safe_font = re.sub(r"[^0-9A-Za-z_]+", "_", east_asia_font)
	style_name = f"GovPara_{safe_font}_{size_pt}"
	try:
		style = doc.styles[style_name]
	except KeyError:
		style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

	_set_style_font(style=style, east_asia_font=east_asia_font, size_pt=size_pt)
	style_fmt = style.paragraph_format
	style_fmt.space_before = Pt(0)
	style_fmt.space_after = Pt(0)
	style_fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
	style_fmt.line_spacing = Pt(LINE_SPACING_PT)
	style_fmt.first_line_indent = Pt(0)
	return style


def _set_paragraph_layout(
	paragraph: Any,
	first_line_indent: bool = False,
	align_center: bool = False,
	align_right: bool = False,
) -> None:
	fmt = paragraph.paragraph_format
	fmt.space_before = Pt(0)
	fmt.space_after = Pt(0)
	fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
	fmt.line_spacing = Pt(LINE_SPACING_PT)
	fmt.first_line_indent = Pt(BODY_FIRST_LINE_INDENT_PT) if first_line_indent else Pt(0)
	if align_center:
		paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
	elif align_right:
		paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	else:
		paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_styled_paragraph(
	doc: Document,
	text: str,
	east_asia_font: str,
	size_pt: int,
	first_line_indent: bool = False,
	align_center: bool = False,
	align_right: bool = False,
	bold: bool = False,
) -> Any:
	style = _get_or_create_paragraph_style(doc=doc, east_asia_font=east_asia_font, size_pt=size_pt)
	paragraph = doc.add_paragraph(style=style)
	_set_paragraph_layout(paragraph, first_line_indent=first_line_indent, align_center=align_center, align_right=align_right)
	run = paragraph.add_run(text)
	_set_run_font(run, east_asia_font=east_asia_font, size_pt=size_pt, bold=bold)
	return paragraph


def _add_warning_paragraph(doc: Document, text: str) -> None:
	paragraph = doc.add_paragraph()
	_set_paragraph_layout(paragraph, first_line_indent=True)
	run = paragraph.add_run(text)
	_set_run_font(run, east_asia_font=BODY_FONT, size_pt=SIZE_PT_LEVEL3, bold=True)
	run.font.color.rgb = COLOR_RED


def _extract_tag_value(summary: str, tag: str) -> str:
	text = str(summary).strip()
	if not text:
		return ""
	pattern = rf"【{re.escape(tag)}】[：:]?\s*(.*?)(?=【[^】]+】|$)"
	match = re.search(pattern, text)
	if not match:
		return ""
	value = match.group(1).strip()
	value = re.sub(r"\s+", " ", value)
	return value


def _extract_core_summary(summary: str) -> str:
	core = _extract_tag_value(summary, "核心主旨")
	if core:
		return core

	first_line = _summary_lines(summary)[0]
	first_line = first_line.replace("【核心主旨】：", "").replace("【核心主旨】:", "").strip()
	if "【关键动作】" in first_line:
		first_line = first_line.split("【关键动作】", maxsplit=1)[0].strip()
	if "【涉及范围】" in first_line:
		first_line = first_line.split("【涉及范围】", maxsplit=1)[0].strip()
	return first_line or "未提及"


def _extract_involved_people(data: Mapping[str, Any], summary: str) -> str:
	candidates = [
		data.get("involved_people"),
		data.get("target_people"),
		data.get("target_audience"),
		data.get("coverage"),
		data.get("scope"),
	]
	for value in candidates:
		text = str(value).strip() if value is not None else ""
		if text:
			return text

	scope_from_summary = _extract_tag_value(summary, "涉及范围")
	if scope_from_summary:
		return scope_from_summary
	return "未提及"


def _split_by_semicolon(text: str) -> list[str]:
	raw = str(text).strip()
	if not raw:
		return []
	normalized = raw.replace(";", "；")
	parts = [part.strip() for part in normalized.split("；")]
	return [part for part in parts if part]


def _add_semicolon_split_paragraphs(
	doc: Document,
	text: str,
	east_asia_font: str,
	size_pt: int,
	first_line_indent: bool,
) -> None:
	parts = _split_by_semicolon(text)
	if not parts:
		_add_styled_paragraph(
			doc,
			text=_safe_text(text, fallback="未提及"),
			east_asia_font=east_asia_font,
			size_pt=size_pt,
			first_line_indent=first_line_indent,
		)
		return
	for part in parts:
		_add_styled_paragraph(
			doc,
			text=part,
			east_asia_font=east_asia_font,
			size_pt=size_pt,
			first_line_indent=first_line_indent,
		)


def _resolve_file_name(data: Mapping[str, Any], context: Mapping[str, Any]) -> str:
	source_file_raw = _safe_text(data.get("source_file"), fallback="")
	if source_file_raw:
		return Path(source_file_raw).name
	return _safe_text(data.get("doc_id"), fallback=_safe_text(context.get("title"), fallback="未命名文件"))


def _append_word_document_sections(doc: Document, data: Mapping[str, Any], section_index: int) -> None:
	context = _build_context(data)
	file_name = _resolve_file_name(data=data, context=context)
	summary_text = _safe_text(context.get("summary"), fallback="未提及")
	core_summary = _extract_core_summary(summary_text)
	involved_people = _extract_involved_people(data=data, summary=summary_text)
	warning_text = _safe_text(context.get("warning"), fallback="")
	crawl_error_url = _safe_text(context.get("crawl_error_url"), fallback="")
	crawl_error_reason = _safe_text(context.get("crawl_error_reason"), fallback="")

	_add_styled_paragraph(
		doc,
		text=f"{_chinese_index(max(1, section_index))}、{file_name}",
		east_asia_font=HEADING2_FONT,
		size_pt=SIZE_PT_LEVEL3,
	)

	if warning_text:
		_add_styled_paragraph(doc, text="0、抓取告警", east_asia_font=HEADING3_FONT, size_pt=SIZE_PT_LEVEL3)
		_add_warning_paragraph(doc, warning_text)
		if crawl_error_url:
			_add_warning_paragraph(doc, f"原链接：{crawl_error_url}")
		if crawl_error_reason:
			_add_warning_paragraph(doc, f"失败原因：{crawl_error_reason}")

	_add_styled_paragraph(doc, text="1、核心主旨", east_asia_font=HEADING3_FONT, size_pt=SIZE_PT_LEVEL3)
	_add_styled_paragraph(
		doc,
		text=_safe_text(core_summary, fallback="未提及"),
		east_asia_font=BODY_FONT,
		size_pt=SIZE_PT_LEVEL3,
		first_line_indent=True,
	)

	tasks = context.get("tasks", []) if isinstance(context.get("tasks"), list) else []
	_add_styled_paragraph(doc, text="2、关键动作", east_asia_font=HEADING3_FONT, size_pt=SIZE_PT_LEVEL3)
	if tasks:
		for idx, task in enumerate(tasks[:8], start=1):
			if not isinstance(task, Mapping):
				continue
			name = _safe_text(task.get("task_name"), fallback="未提及")
			owner = _safe_text(task.get("owner"), fallback="相关责任部门")
			deadline = _safe_text(task.get("deadline_display"), fallback=_safe_text(task.get("deadline"), fallback="未提及"))
			_add_styled_paragraph(
				doc,
				text=f"{idx}. {name}（责任：{owner}；时限：{deadline}）",
				east_asia_font=BODY_FONT,
				size_pt=SIZE_PT_LEVEL3,
				first_line_indent=True,
			)
	else:
		_add_styled_paragraph(doc, text="未提取到关键动作。", east_asia_font=BODY_FONT, size_pt=SIZE_PT_LEVEL3, first_line_indent=True)

	_add_styled_paragraph(doc, text="3、设计范围", east_asia_font=HEADING3_FONT, size_pt=SIZE_PT_LEVEL3)
	design_scope_lines = [
		f"文种：{_safe_text(context.get('doc_type'), fallback='未提及')}",
		f"发文单位：{_safe_text(context.get('issuing_department'), fallback='未提及')}",
		f"发布日期：{_safe_text(context.get('publish_date'), fallback='未提及')}",
		f"文号：{_safe_text(context.get('document_no'), fallback='未提及')}",
		f"涉及人群：{_safe_text(involved_people, fallback='未提及')}",
	]
	for line in design_scope_lines:
		_add_styled_paragraph(doc, text=line, east_asia_font=BODY_FONT, size_pt=SIZE_PT_LEVEL3, first_line_indent=True)

	_add_styled_paragraph(doc, text="4、任务清单", east_asia_font=HEADING3_FONT, size_pt=SIZE_PT_LEVEL3)
	if tasks:
		for idx, task in enumerate(tasks, start=1):
			if not isinstance(task, Mapping):
				continue
			task_name = _safe_text(task.get("task_name"), fallback="未提及")
			owner = _safe_text(task.get("owner"), fallback="相关责任部门")
			deadline = _safe_text(task.get("deadline_display"), fallback=_safe_text(task.get("deadline"), fallback="未提及"))
			urgency_label = _safe_text(task.get("urgency", {}).get("label"), fallback="未提及") if isinstance(task.get("urgency"), Mapping) else "未提及"
			deliverables = task.get("deliverables", []) if isinstance(task.get("deliverables"), list) else []
			deliverables_text = "；".join(str(item).strip() for item in deliverables if str(item).strip()) or "未提及"

			_add_styled_paragraph(
				doc,
				text=f"（{idx}）任务名称：{task_name}",
				east_asia_font=BODY_FONT,
				size_pt=SIZE_PT_LEVEL3,
				first_line_indent=False,
				bold=True,
			)
			_add_semicolon_split_paragraphs(
				doc=doc,
				text=f"责任部门：{owner}；截止时间：{deadline}；紧急度：{urgency_label}",
				east_asia_font=BODY_FONT,
				size_pt=SIZE_PT_LEVEL3,
				first_line_indent=True,
			)
			_add_semicolon_split_paragraphs(
				doc=doc,
				text=f"交付物：{deliverables_text}",
				east_asia_font=BODY_FONT,
				size_pt=SIZE_PT_LEVEL3,
				first_line_indent=True,
			)
	else:
		_add_styled_paragraph(doc, text="未提取到任务清单。", east_asia_font=BODY_FONT, size_pt=SIZE_PT_LEVEL3, first_line_indent=True)

	risks = context.get("risks_or_unclear_points", []) if isinstance(context.get("risks_or_unclear_points"), list) else []
	if risks:
		_add_styled_paragraph(doc, text="5、风险提示", east_asia_font=HEADING3_FONT, size_pt=SIZE_PT_LEVEL3)
		for line in risks:
			_add_styled_paragraph(
				doc,
				text=_safe_text(line, fallback="未提及"),
				east_asia_font=BODY_FONT,
				size_pt=SIZE_PT_LEVEL3,
				first_line_indent=True,
			)

	questions = context.get("follow_up_questions", []) if isinstance(context.get("follow_up_questions"), list) else []
	if questions:
		_add_styled_paragraph(doc, text="6、待确认事项", east_asia_font=HEADING3_FONT, size_pt=SIZE_PT_LEVEL3)
		for line in questions:
			_add_styled_paragraph(
				doc,
				text=_safe_text(line, fallback="未提及"),
				east_asia_font=BODY_FONT,
				size_pt=SIZE_PT_LEVEL3,
				first_line_indent=True,
			)


def _render_word_bundle_report(data_items: list[Mapping[str, Any]], output_path: Path) -> Path:
	if not data_items:
		raise ValueError("No input data for bundle word report")

	doc = Document()
	model_display = _resolve_model_display(data_items[0])

	_add_styled_paragraph(
		doc,
		text=TITLE_TEXT,
		east_asia_font=TITLE_FONT,
		size_pt=SIZE_PT_LEVEL1,
		align_center=True,
	)
	_add_styled_paragraph(
		doc,
		text=f"--公文速阅智能体基于【{model_display}】生成的报告",
		east_asia_font=SUBTITLE_FONT,
		size_pt=SIZE_PT_LEVEL2,
		align_right=True,
	)
	_add_styled_paragraph(
		doc,
		text=_build_word_timestamp_line(),
		east_asia_font=SUBTITLE_FONT,
		size_pt=SIZE_PT_LEVEL2,
		align_center=True,
	)
	_add_styled_paragraph(doc, text="", east_asia_font=SUBTITLE_FONT, size_pt=SIZE_PT_LEVEL2)

	for idx, data in enumerate(data_items, start=1):
		_append_word_document_sections(doc=doc, data=data, section_index=idx)

	output_path.parent.mkdir(parents=True, exist_ok=True)
	doc.save(output_path)
	log_step(
		LOGGER,
		"report.render",
		"ReportRenderer",
		"RenderWordBundleDone",
		f"items={len(data_items)} path={to_relative_path(output_path)}",
	)
	return output_path


def _parse_deadline_date(value: str) -> date | None:
	text = str(value).strip()
	full = re.match(r"^(\d{4})-(\d{2})-(\d{2})$", text)
	if full:
		year, month, day = map(int, full.groups())
		try:
			return date(year, month, day)
		except ValueError:
			return None

	month_only = re.match(r"^(\d{4})-(\d{2})$", text)
	if month_only:
		year, month = map(int, month_only.groups())
		try:
			return date(year, month, 1)
		except ValueError:
			return None

	only_year = re.match(r"^(\d{4})$", text)
	if only_year:
		year = int(only_year.group(1))
		return date(year, 1, 1)

	return None


def _extract_first_iso_date(value: str) -> date | None:
	text = str(value).strip()
	match = re.search(r"(\d{4}-\d{2}-\d{2})", text)
	if not match:
		return None
	return _parse_deadline_date(match.group(1))


def _task_id_sort_key(task_id: str) -> tuple[int, str]:
	value = str(task_id).strip()
	match = re.search(r"(\d+)", value)
	if match:
		return int(match.group(1)), value
	return 10**9, value


def _is_long_term_task(deadline_display: str) -> bool:
	value = str(deadline_display).strip()
	if value in {"长期有效", "按需执行", "无", "未提及", ""}:
		return True
	periodic_keywords = ["每年", "每月", "每周", "每季度", "每学期", "定期", "常态"]
	return any(keyword in value for keyword in periodic_keywords)


def _normalize_urgency(task: Mapping[str, Any], deadline_display: str, deadline_start: str) -> dict[str, Any]:
	raw_urgency = task.get("urgency", {}) if isinstance(task.get("urgency"), Mapping) else {}

	try:
		raw_score = int(raw_urgency.get("score", 0))
	except Exception:  # pylint: disable=broad-except
		raw_score = 0

	raw_color = _safe_text(raw_urgency.get("color"), fallback="").lower()
	raw_level = _safe_text(raw_urgency.get("level"), fallback="").upper()
	raw_label = _safe_text(raw_urgency.get("label"), fallback="")

	effective_start = _parse_deadline_date(deadline_start) or _parse_deadline_date(deadline_display) or _extract_first_iso_date(deadline_display)

	if effective_start is None and _is_long_term_task(deadline_display):
		color = "gray"
		level = "NONE"
		score = 0
		label = "长效制度"
	else:
		if raw_color not in {"red", "yellow", "green", "gray"}:
			if raw_level == "HIGH" or raw_score >= 70:
				raw_color = "red"
			elif raw_level == "MEDIUM" or raw_score >= 40:
				raw_color = "yellow"
			elif raw_level == "NONE":
				raw_color = "gray"
			else:
				raw_color = "green"

		color = raw_color
		if color == "red":
			level = "HIGH"
			label = raw_label if raw_label and raw_label != "未提及" else "特急"
			score = max(raw_score, 70)
		elif color == "yellow":
			level = "MEDIUM"
			label = raw_label if raw_label and raw_label != "未提及" else "紧急"
			score = max(raw_score, 40)
		elif color == "gray":
			level = "NONE"
			label = raw_label if raw_label and raw_label != "未提及" else "长效制度"
			score = 0
		else:
			level = "LOW"
			label = raw_label if raw_label and raw_label != "未提及" else "常规"
			score = max(0, raw_score)

	emoji_map = {
		"red": "🔴",
		"yellow": "🟡",
		"green": "🟢",
		"gray": "⚪",
	}

	return {
		"score": score,
		"level": level,
		"color": color,
		"label": label,
		"emoji": emoji_map.get(color, "⚪"),
	}


def _normalize_task(
	task: Mapping[str, Any],
	index: int,
	event_by_task: Mapping[str, Mapping[str, Any]],
) -> dict[str, Any]:
	task_id = _safe_text(task.get("task_id"), fallback=f"task_{index:03d}")
	deadline_display = _safe_text(task.get("deadline_display"), fallback="")
	if deadline_display in {"无", "未提及", ""}:
		deadline_display = _safe_text(task.get("deadline"), fallback="未提及")

	deliverables_raw = task.get("deliverables", [])
	deliverables = [str(item).strip() for item in deliverables_raw] if isinstance(deliverables_raw, list) else []
	deliverables = [item for item in deliverables if item]

	deadline_start = _safe_text(task.get("deadline_start"), fallback="")
	if deadline_start in {"无", "未提及"}:
		deadline_start = ""

	event = event_by_task.get(task_id, {}) if isinstance(event_by_task, Mapping) else {}
	calendar_uid = _safe_text(event.get("uid"), fallback="") if isinstance(event, Mapping) else ""

	source_anchor_raw = task.get("source_anchor", {}) if isinstance(task.get("source_anchor"), Mapping) else {}
	source_anchor = {
		"block_id": _safe_text(source_anchor_raw.get("block_id"), fallback="未提及"),
		"quote": _safe_text(source_anchor_raw.get("quote"), fallback="未提及"),
	}

	urgency = _normalize_urgency(task, deadline_display=deadline_display, deadline_start=deadline_start)

	return {
		"task_id": task_id,
		"task_name": _safe_text(task.get("task_name")),
		"owner": _safe_text(task.get("owner"), fallback="相关责任部门"),
		"deadline": _safe_text(task.get("deadline"), fallback="未提及"),
		"deadline_display": deadline_display,
		"deadline_start": deadline_start,
		"action_suggestion": _safe_text(task.get("action_suggestion"), fallback="按要求执行"),
		"deliverables": deliverables,
		"source_anchor": source_anchor,
		"calendar_uid": calendar_uid,
		"urgency": urgency,
	}


def _sort_tasks(tasks: list[dict[str, Any]]) -> list[dict[str, Any]]:
	return sorted(
		tasks,
		key=lambda task: (
			-int(task.get("urgency", {}).get("score", 0)),
			_parse_deadline_date(task.get("deadline_start", ""))
			or _parse_deadline_date(task.get("deadline_display", ""))
			or _extract_first_iso_date(task.get("deadline_display", ""))
			or date.max,
			_task_id_sort_key(task.get("task_id", "")),
		),
	)


def _build_context(data: Mapping[str, Any]) -> dict[str, Any]:
	log_step(
		LOGGER,
		"report.context",
		"ReportRenderer",
		"BuildContextStart",
		f"doc_id={data.get('doc_id', 'unknown')} raw_tasks={len(data.get('tasks', [])) if isinstance(data.get('tasks'), list) else 0}",
	)
	calendar_raw = data.get("calendar", {}) if isinstance(data.get("calendar"), Mapping) else {}
	events_raw = calendar_raw.get("events", [])
	events = events_raw if isinstance(events_raw, list) else []
	event_by_task: dict[str, Mapping[str, Any]] = {}
	for item in events:
		if not isinstance(item, Mapping):
			continue
		task_id = str(item.get("task_id", "")).strip()
		if not task_id:
			continue
		event_by_task[task_id] = item

	ics_file_raw = str(calendar_raw.get("ics_file", "")).strip()
	ics_attachment_name = Path(ics_file_raw).name if ics_file_raw else f"{_safe_text(data.get('doc_id'), fallback='document')}.ics"

	raw_tasks = data.get("tasks", [])
	tasks = raw_tasks if isinstance(raw_tasks, list) else []
	normalized_tasks = [
		_normalize_task(task, index=i, event_by_task=event_by_task)
		for i, task in enumerate(tasks, start=1)
		if isinstance(task, Mapping)
	]
	tasks_sorted = _sort_tasks(normalized_tasks)

	risks_raw = data.get("risks_or_unclear_points", [])
	risks = [str(item).strip() for item in risks_raw] if isinstance(risks_raw, list) else []
	risks = [item for item in risks if item]

	questions_raw = data.get("follow_up_questions", [])
	questions = [str(item).strip() for item in questions_raw] if isinstance(questions_raw, list) else []
	questions = [item for item in questions if item]

	summary = _safe_text(data.get("summary"), fallback="未提及")
	status = _safe_text(data.get("status"), fallback="success").lower()
	crawl_error = data.get("crawl_error", {}) if isinstance(data.get("crawl_error"), Mapping) else {}
	crawl_error_url = _safe_text(crawl_error.get("url"), fallback=_safe_text(data.get("source_url"), fallback=""))
	crawl_error_reason = _safe_text(crawl_error.get("reason"), fallback=_safe_text(data.get("reason"), fallback=""))
	warning = _safe_text(data.get("warning"), fallback="")
	if not warning and status == "failed":
		warning = "警告：该网址来源抓取失败，请人工核实原链接。"

	return {
		"title": _safe_text(data.get("title"), fallback="公文报告"),
		"status": status,
		"document_no": _safe_text(data.get("document_no"), fallback="未提及"),
		"publish_date": _safe_text(data.get("publish_date"), fallback="未提及"),
		"issuing_department": _safe_text(data.get("issuing_department"), fallback="未提及"),
		"doc_type": _safe_text(data.get("doc_type"), fallback="未提及"),
		"summary": summary,
		"summary_lines": _summary_lines(summary),
		"warning": warning,
		"crawl_error_url": crawl_error_url,
		"crawl_error_reason": crawl_error_reason,
		"tasks": tasks_sorted,
		"tasks_html": tasks_sorted,
		"calendar_attachment_name": ics_attachment_name,
		"calendar_hint": "请在邮件附件中打开 .ics 文件并导入日历。",
		"risks_or_unclear_points": risks,
		"follow_up_questions": questions,
	}



def _build_env(template_dir: Path) -> Environment:
	return Environment(
		loader=FileSystemLoader(str(template_dir)),
		autoescape=select_autoescape(["html", "xml"]),
		trim_blocks=True,
		lstrip_blocks=True,
	)


def _normalize_report_formats(formats: Iterable[str] | None) -> list[str]:
	if formats is None:
		return ["md", "html", "docx"]
	seen: list[str] = []
	for item in formats:
		key = str(item or "").strip().lower()
		if key not in SUPPORTED_REPORT_FORMATS:
			continue
		if key in seen:
			continue
		seen.append(key)
	return seen or ["md", "html", "docx"]


def _render_markdown_content(data: Mapping[str, Any], template_dir: Path) -> str:
	env = _build_env(template_dir)
	template = env.get_template("report_template.md")
	context = _build_context(data)
	return template.render(**context)


def _render_html_content(data: Mapping[str, Any], template_dir: Path) -> str:
	env = _build_env(template_dir)
	template = env.get_template("report_template.html")
	context = _build_context(data)
	return template.render(**context)


def _render_markdown_bundle_report(data_items: list[Mapping[str, Any]], output_path: Path) -> Path:
	if not data_items:
		raise ValueError("No input data for bundle markdown report")

	timestamp_line = _build_word_timestamp_line()
	lines: list[str] = [
		"# 公文速阅报告报送",
		f"> {timestamp_line}",
		"",
	]

	for index, data in enumerate(data_items, start=1):
		context = _build_context(data)
		lines.extend(
			[
				"---",
				"",
				f"## {_chinese_index(index)}、{_safe_text(context.get('title'), fallback='未命名公文')}",
				"",
				f"- 发文部门：{_safe_text(context.get('issuing_department'), fallback='未提及')}",
				f"- 文号：{_safe_text(context.get('document_no'), fallback='未提及')}",
				f"- 发布日期：{_safe_text(context.get('publish_date'), fallback='未提及')}",
				f"- 公文类型：{_safe_text(context.get('doc_type'), fallback='未提及')}",
				"",
				"### 核心摘要",
			]
		)
		for summary_line in context.get("summary_lines", []):
			lines.append(f"> {_safe_text(summary_line, fallback='未提及')}")

		lines.extend(["", "### 待办清单", ""])
		tasks = context.get("tasks", []) if isinstance(context.get("tasks"), list) else []
		if tasks:
			for task in tasks:
				if not isinstance(task, Mapping):
					continue
				lines.append(f"- [ ] **{_safe_text(task.get('task_name'), fallback='未提及')}**")
				lines.append(f"  - 责任人：{_safe_text(task.get('owner'), fallback='相关责任部门')}")
				lines.append(f"  - 截止：{_safe_text(task.get('deadline_display'), fallback='未提及')}")
		else:
			lines.append("- [ ] 未提取到任务")

		lines.extend(["", "### 风险与追问", ""])
		risks = context.get("risks_or_unclear_points", []) if isinstance(context.get("risks_or_unclear_points"), list) else []
		questions = context.get("follow_up_questions", []) if isinstance(context.get("follow_up_questions"), list) else []
		if risks:
			for risk in risks:
				lines.append(f"- 风险：{_safe_text(risk, fallback='未提及')}")
		else:
			lines.append("- 风险：暂未识别到明显风险")
		if questions:
			for question in questions:
				lines.append(f"- 追问：{_safe_text(question, fallback='未提及')}")
		else:
			lines.append("- 追问：暂无补充追问")

	content = "\n".join(lines).strip() + "\n"
	output_path.parent.mkdir(parents=True, exist_ok=True)
	output_path.write_text(content, encoding="utf-8")
	log_step(
		LOGGER,
		"report.render",
		"ReportRenderer",
		"RenderMarkdownBundleDone",
		f"items={len(data_items)} path={to_relative_path(output_path)}",
	)
	return output_path


def _render_html_bundle_report(data_items: list[Mapping[str, Any]], output_path: Path) -> Path:
	if not data_items:
		raise ValueError("No input data for bundle html report")

	sections: list[str] = []
	for index, data in enumerate(data_items, start=1):
		context = _build_context(data)
		tasks = context.get("tasks", []) if isinstance(context.get("tasks"), list) else []
		task_rows = "".join(
			f"<li><strong>{html_escape(_safe_text(task.get('task_name'), fallback='未提及'))}</strong>"
			f"<span>责任：{html_escape(_safe_text(task.get('owner'), fallback='相关责任部门'))}</span>"
			f"<span>截止：{html_escape(_safe_text(task.get('deadline_display'), fallback='未提及'))}</span></li>"
			for task in tasks
			if isinstance(task, Mapping)
		)
		if not task_rows:
			task_rows = "<li><strong>未提取到任务</strong></li>"

		sections.append(
			f"""
			<section class=\"doc-card\">
			  <h2>{_chinese_index(index)}、{html_escape(_safe_text(context.get('title'), fallback='未命名公文'))}</h2>
			  <p class=\"meta\">发文部门：{html_escape(_safe_text(context.get('issuing_department'), fallback='未提及'))} ｜ 文号：{html_escape(_safe_text(context.get('document_no'), fallback='未提及'))} ｜ 发布日期：{html_escape(_safe_text(context.get('publish_date'), fallback='未提及'))}</p>
			  <div class=\"summary\">{html_escape(_safe_text(context.get('summary'), fallback='未提及'))}</div>
			  <h3>任务清单</h3>
			  <ul class=\"task-list\">{task_rows}</ul>
			</section>
			"""
		)

	html_content = f"""<!doctype html>
<html lang=\"zh-CN\">
<head>
  <meta charset=\"utf-8\" />
  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\" />
  <title>公文速阅报告汇总</title>
  <style>
    body {{ margin: 0; padding: 24px; font-family: "Microsoft YaHei", "PingFang SC", sans-serif; background: #f4f7fb; color: #1f2937; }}
    .container {{ max-width: 1180px; margin: 0 auto; }}
    .header {{ background: #fff; border: 1px solid #e5eaf0; border-radius: 12px; padding: 18px 20px; margin-bottom: 14px; }}
    .header h1 {{ margin: 0; color: #1b2a41; font-size: 30px; }}
    .header p {{ margin: 8px 0 0; color: #53657d; }}
    .doc-card {{ background: #fff; border: 1px solid #e5eaf0; border-left: 5px solid #b0182e; border-radius: 12px; padding: 16px 18px; margin-bottom: 12px; }}
    .doc-card h2 {{ margin: 0 0 6px; color: #1f3552; font-size: 22px; }}
    .meta {{ margin: 0 0 10px; color: #61748c; font-size: 13px; }}
    .summary {{ margin: 0 0 12px; background: #fff6f7; border-radius: 8px; padding: 10px 12px; line-height: 1.7; }}
    .task-list {{ margin: 0; padding-left: 18px; }}
    .task-list li {{ margin: 6px 0; }}
    .task-list span {{ display: inline-block; margin-left: 10px; color: #4f647e; font-size: 13px; }}
  </style>
</head>
<body>
  <div class=\"container\">
    <section class=\"header\">
      <h1>公文速阅报告报送（汇总版）</h1>
      <p>{html_escape(_build_word_timestamp_line())}</p>
    </section>
    {''.join(sections)}
  </div>
</body>
</html>
"""

	output_path.parent.mkdir(parents=True, exist_ok=True)
	output_path.write_text(html_content, encoding="utf-8")
	log_step(
		LOGGER,
		"report.render",
		"ReportRenderer",
		"RenderHtmlBundleDone",
		f"items={len(data_items)} path={to_relative_path(output_path)}",
	)
	return output_path


def render_bundle_reports(
	data_items: list[Mapping[str, Any]],
	output_dir: Path = DEFAULT_REPORT_DIR,
	template_dir: Path = DEFAULT_TEMPLATE_DIR,
	bundle_formats: Iterable[str] | None = None,
) -> dict[str, Path]:
	formats = _normalize_report_formats(bundle_formats)
	if not data_items:
		return {}

	base_name = _safe_filename(_build_word_report_filename())
	result: dict[str, Path] = {}

	if "md" in formats:
		md_path = output_dir / f"{base_name}.md"
		result["md"] = _render_markdown_bundle_report(data_items=data_items, output_path=md_path)

	if "html" in formats:
		html_path = output_dir / f"{base_name}.html"
		result["html"] = _render_html_bundle_report(data_items=data_items, output_path=html_path)

	if "docx" in formats:
		docx_path = output_dir / f"{base_name}.docx"
		result["docx"] = _render_word_bundle_report(data_items=data_items, output_path=docx_path)

	return result


def render_markdown_report(
	data: Mapping[str, Any],
	output_path: Path,
	template_dir: Path = DEFAULT_TEMPLATE_DIR,
) -> Path:
	env = _build_env(template_dir)
	template = env.get_template("report_template.md")
	context = _build_context(data)
	content = template.render(**context)
	output_path.parent.mkdir(parents=True, exist_ok=True)
	output_path.write_text(content, encoding="utf-8")
	log_step(
		LOGGER,
		"report.render",
		"ReportRenderer",
		"RenderMarkdownDone",
		f"doc_id={context.get('title', 'unknown')} path={to_relative_path(output_path)}",
	)
	return output_path


def render_html_report(
	data: Mapping[str, Any],
	output_path: Path,
	template_dir: Path = DEFAULT_TEMPLATE_DIR,
) -> Path:
	env = _build_env(template_dir)
	template = env.get_template("report_template.html")
	context = _build_context(data)
	content = template.render(**context)
	output_path.parent.mkdir(parents=True, exist_ok=True)
	output_path.write_text(content, encoding="utf-8")
	log_step(
		LOGGER,
		"report.render",
		"ReportRenderer",
		"RenderHTMLDone",
		f"doc_id={context.get('title', 'unknown')} path={to_relative_path(output_path)}",
	)
	return output_path


def render_word_report(
	data: Mapping[str, Any],
	output_path: Path,
	section_index: int = 1,
) -> Path:
	context = _build_context(data)
	model_display = _resolve_model_display(data)

	doc = Document()

	_add_styled_paragraph(
		doc,
		text=TITLE_TEXT,
		east_asia_font=TITLE_FONT,
		size_pt=SIZE_PT_LEVEL1,
		align_center=True,
	)
	_add_styled_paragraph(
		doc,
		text=f"--公文速阅智能体基于【{model_display}】生成的报告",
		east_asia_font=SUBTITLE_FONT,
		size_pt=SIZE_PT_LEVEL2,
		align_right=True,
	)
	_add_styled_paragraph(
		doc,
		text=_build_word_timestamp_line(),
		east_asia_font=SUBTITLE_FONT,
		size_pt=SIZE_PT_LEVEL2,
		align_center=True,
	)
	_add_styled_paragraph(doc, text="", east_asia_font=SUBTITLE_FONT, size_pt=SIZE_PT_LEVEL2)
	_append_word_document_sections(doc=doc, data=data, section_index=section_index)

	output_path.parent.mkdir(parents=True, exist_ok=True)
	doc.save(output_path)
	log_step(
		LOGGER,
		"report.render",
		"ReportRenderer",
		"RenderWordDone",
		f"doc_id={context.get('title', 'unknown')} path={to_relative_path(output_path)}",
	)
	return output_path


def render_report_from_json(
	json_path: Path,
	output_dir: Path = DEFAULT_REPORT_DIR,
	template_dir: Path = DEFAULT_TEMPLATE_DIR,
	section_index: int = 1,
) -> tuple[Path, Path, Path]:
	rendered = render_selected_reports_from_json(
		json_path=json_path,
		output_dir=output_dir,
		template_dir=template_dir,
		include_formats=["md", "html", "docx"],
		section_index=section_index,
	)
	return rendered["md"], rendered["html"], rendered["docx"]


def render_selected_reports_from_json(
	json_path: Path,
	output_dir: Path = DEFAULT_REPORT_DIR,
	template_dir: Path = DEFAULT_TEMPLATE_DIR,
	include_formats: Iterable[str] | None = None,
	section_index: int = 1,
) -> dict[str, Path]:
	log_step(LOGGER, "report.render", "ReportRenderer", "RenderFromJsonStart", f"json_file={to_relative_path(json_path)}")
	data = json.loads(json_path.read_text(encoding="utf-8"))
	if not isinstance(data, Mapping):
		raise ValueError(f"JSON top-level must be object: {json_path}")
	formats = _normalize_report_formats(include_formats)

	file_stem = _safe_filename(_safe_text(data.get("doc_id"), fallback=json_path.stem))
	rendered_paths: dict[str, Path] = {}

	if "md" in formats:
		md_path = output_dir / f"{file_stem}.report.md"
		render_markdown_report(data=data, output_path=md_path, template_dir=template_dir)
		rendered_paths["md"] = md_path

	if "html" in formats:
		html_path = output_dir / f"{file_stem}.report.html"
		render_html_report(data=data, output_path=html_path, template_dir=template_dir)
		rendered_paths["html"] = html_path

	if "docx" in formats:
		docx_path = output_dir / f"{file_stem}.report.docx"
		render_word_report(data=data, output_path=docx_path, section_index=section_index)
		rendered_paths["docx"] = docx_path

	md_rel = to_relative_path(rendered_paths["md"]) if "md" in rendered_paths else "-"
	html_rel = to_relative_path(rendered_paths["html"]) if "html" in rendered_paths else "-"
	docx_rel = to_relative_path(rendered_paths["docx"]) if "docx" in rendered_paths else "-"
	log_step(
		LOGGER,
		"report.render",
		"ReportRenderer",
		"RenderFromJsonDone",
		f"json_file={to_relative_path(json_path)} md={md_rel} html={html_rel} docx={docx_rel}",
	)
	return rendered_paths


def render_reports_from_cache(
	cache_dir: Path = DEFAULT_CACHE_DIR,
	output_dir: Path = DEFAULT_REPORT_DIR,
	template_dir: Path = DEFAULT_TEMPLATE_DIR,
) -> list[tuple[Path, Path, Path]]:
	log_step(LOGGER, "report.batch", "ReportRenderer", "BatchRenderStart", f"cache_dir={to_relative_path(cache_dir)}")
	json_files = sorted(cache_dir.glob("*.json"))
	results_tmp: list[tuple[Path, Path]] = []
	data_items: list[Mapping[str, Any]] = []

	for json_file in json_files:
		log_step(LOGGER, "report.render", "ReportRenderer", "RenderFromJsonStart", f"json_file={to_relative_path(json_file)}")
		data = json.loads(json_file.read_text(encoding="utf-8"))
		if not isinstance(data, Mapping):
			raise ValueError(f"JSON top-level must be object: {json_file}")
		data_items.append(data)

		file_stem = _safe_filename(_safe_text(data.get("doc_id"), fallback=json_file.stem))
		md_path = output_dir / f"{file_stem}.report.md"
		html_path = output_dir / f"{file_stem}.report.html"

		render_markdown_report(data=data, output_path=md_path, template_dir=template_dir)
		render_html_report(data=data, output_path=html_path, template_dir=template_dir)
		results_tmp.append((md_path, html_path))
		log_step(
			LOGGER,
			"report.render",
			"ReportRenderer",
			"RenderFromJsonDone",
			f"json_file={to_relative_path(json_file)} md={to_relative_path(md_path)} html={to_relative_path(html_path)}",
		)

	if not data_items:
		log_step(LOGGER, "report.batch", "ReportRenderer", "BatchRenderDone", f"files=0 output_dir={to_relative_path(output_dir)}")
		return []

	bundle_title = _build_word_report_filename()
	bundle_docx_path = output_dir / f"{_safe_filename(bundle_title)}.docx"
	_render_word_bundle_report(data_items=data_items, output_path=bundle_docx_path)

	results = [(md_path, html_path, bundle_docx_path) for md_path, html_path in results_tmp]
	log_step(
		LOGGER,
		"report.batch",
		"ReportRenderer",
		"BatchRenderDone",
		f"files={len(results)} bundle_docx={to_relative_path(bundle_docx_path)} output_dir={to_relative_path(output_dir)}",
	)
	return results


def _parse_args() -> argparse.Namespace:
	parser = argparse.ArgumentParser(description="Render markdown/html/word reports from processed JSON cache.")
	parser.add_argument("--json-file", type=str, default="", help="Render a single JSON file.")
	parser.add_argument("--cache-dir", type=str, default=str(DEFAULT_CACHE_DIR), help="Directory containing cached JSON files.")
	parser.add_argument("--output-dir", type=str, default=str(DEFAULT_REPORT_DIR), help="Directory to save rendered reports.")
	parser.add_argument("--template-dir", type=str, default=str(DEFAULT_TEMPLATE_DIR), help="Directory containing Jinja2 templates.")
	return parser.parse_args()


def main() -> None:
	setup_logger(force_reconfigure=True)
	main_logger = get_logger("report_renderer")
	session = get_log_session()
	if session is not None:
		log_step(
			main_logger,
			"startup",
			"ReportRenderer",
			"LoggerReady",
			f"run_id={session.run_id} log_file={to_relative_path(session.log_file)}",
		)
	args = _parse_args()
	template_dir = Path(args.template_dir)
	output_dir = Path(args.output_dir)
	log_step(
		main_logger,
		"startup",
		"ReportRenderer",
		"ArgsParsed",
		f"json_file={to_relative_path(args.json_file) if args.json_file else 'ALL'} cache_dir={to_relative_path(args.cache_dir)} output_dir={to_relative_path(args.output_dir)}",
	)

	if args.json_file:
		md_path, html_path, docx_path = render_report_from_json(
			json_path=Path(args.json_file),
			output_dir=output_dir,
			template_dir=template_dir,
		)
		print(to_relative_path(md_path))
		print(to_relative_path(html_path))
		print(to_relative_path(docx_path))
		log_step(
			main_logger,
			"shutdown",
			"ReportRenderer",
			"SingleRenderCompleted",
			f"md={to_relative_path(md_path)} html={to_relative_path(html_path)} docx={to_relative_path(docx_path)}",
		)
		return

	results = render_reports_from_cache(
		cache_dir=Path(args.cache_dir),
		output_dir=output_dir,
		template_dir=template_dir,
	)
	printed_docx: set[Path] = set()
	for md_path, html_path, docx_path in results:
		print(to_relative_path(md_path))
		print(to_relative_path(html_path))
		if docx_path not in printed_docx:
			print(to_relative_path(docx_path))
			printed_docx.add(docx_path)
	log_step(main_logger, "shutdown", "ReportRenderer", "BatchRenderCompleted", f"files={len(results)}")


if __name__ == "__main__":
	main()
